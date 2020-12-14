from docx import Document
from docx.shared import Inches
import docx
# Создаем пустой вордовский документ
document = Document()
"""
Так как все в ворде это параграфы добавляем параграф для примера
"""
# Добавление заголовка

document.add_heading('Во славу Омнисиии!!!')
# Присваиваем ссылку чтобы можно было добавлять параграфы сверху или снизу
name_copp = document.add_paragraph('Центр опережающей профессиональной подготовки Республики Бурятия')
# Вставляем параграф сверху
above_paragraph = name_copp.insert_paragraph_before('Во славу Знания!!!')
# Добавление заголовка с желаемым уровнем
document.add_heading('Это заголовок 4 уровня',level=4)

# Разрыв страницы
document.add_page_break()

document.add_heading('Мысль дня',level=1)
document.add_paragraph('Вера мой щит, разум мой меч.')


# Добавление таблицы

table = document.add_table(rows=5, cols=2)
cell_anchor = table.cell(0,1)

# Записываем в ячейку текст
cell_anchor.text = 'Cassandra Cilian'
document.add_picture('data/Lindy Booth.jpg',width=Inches(4.25))
row = table.rows[1]
row.cells[0].text = 'Foo bar to you.'
row.cells[1].text = 'And a hearty foo bar to you too sir!'

document.save('Prototype.docx')
