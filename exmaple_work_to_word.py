import docx
import os

# Список для сохранения имен документов

# Получаем текущую папку и добавляем еще одну подпапку
folder = os.getcwd() + '\data'


def get_list_files(path, extension='docx'):
    """
    Функция для получения списка всех  файлов внутри определенной папки с определенным разрешением
    :param path:Путь к папке
    :return:список файлов
    """
    paths = []
    # Рекурсивно перебираем папки и файлы находящиеся в папке path
    for root, dirs, files in os.walk(path):
        # Перебираем найденные файлы чтобы отобрать файлы с нужным расширением
        for file in files:
            if file.endswith(extension) and not file.startswith('~'):
                paths.append(os.path.join(root,file))
    return paths

files = get_list_files(folder,'pdf')
print(files)
# Свойства полученных документов
# Перебираем документы и извлекаем свойства

# for file in paths:
#     doc =  docx.Document(file)
#     properties = doc.core_properties
#     print('Автор документа:', properties.author)
#     print('Автор последней правки:', properties.last_modified_by)
#     print('Дата создания документа:', properties.created)
#     print('Дата последней правки:', properties.modified)
#     print('Дата последней печати:', properties.last_printed)
#     print('Количество сохранений:', properties.revision)
#     print('**********************')

"""
Объект Document, представляющий собой весь документ
Список объектов Paragraph – абзацы документа
Список объектов Run – фрагменты текста с различными стилями форматирования (курсив, цвет шрифта и т.п.)
Список объектов Table – таблицы документа
Список объектов Row – строки таблицы
Список объектов Cell – ячейки в строке
Список объектов Column – столбцы таблицы
Список объектов Cell – ячейки в столбце
Список объектов InlineShape – иллюстрации документа
"""

# # Получаем весь текст из документа
# for file in paths:
#     doc =  docx.Document(file)
#     # Создаем список который будет содержать в себе тексты документов
#     text = []
#     for paragraph in doc.paragraphs:
#         text.append(paragraph.text)
#     print('/n'.join(text))

# Извлечение текста выделенного определенным форматом
# for file in paths:
#     doc =  docx.Document(file)
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             if run.bold:
#                 print(run.text)

# Еще стили форматирования
# for file in paths:
#     doc =  docx.Document(file)
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             print('Полужирный текст:', run.bold)
#             print('Подчёркнутый текст:', run.underline)
#             print('Зачёркнутый текст:', run.strike)
#             print('Название шрифта:', run.font.name)
#             print('Цвет текста, RGB:', run.font.color.rgb)
#             print('Цвет заливки текста:', run.font.highlight_color)